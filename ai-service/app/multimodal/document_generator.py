# -*- coding: utf-8 -*-
"""
文档生成模块

提供Word和PDF文档生成功能，支持健康报告等多种模板。

依赖库:
    - python-docx: Word文档生成
    - markdown: Markdown转HTML
    - weasyprint: HTML转PDF (可选)
    - docx2pdf: Word转PDF (Windows平台)

作者: AI Assistant
创建日期: 2026-03-25
"""

import os
import io
import logging
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List, Union
from pathlib import Path

# 配置日志记录
logger = logging.getLogger(__name__)


class DocumentGenerator:
    """
    文档生成器类
    
    支持生成Word和PDF格式的文档，提供多种模板（健康报告、体检报告等）。
    
    Attributes:
        output_dir (str): 文档输出目录
        templates (Dict): 可用模板配置
    """
    
    # 支持的文档类型
    SUPPORTED_DOC_TYPES = ["word", "pdf"]
    
    # 默认模板配置
    DEFAULT_TEMPLATES = {
        "health_report": {
            "name": "健康报告",
            "description": "个人健康评估报告模板",
            "sections": [
                "basic_info",      # 基本信息
                "health_summary",  # 健康摘要
                "vital_signs",     # 生命体征
                "lab_results",     # 检验结果
                "recommendations", # 健康建议
                "follow_up"        # 随访计划
            ]
        },
        "medical_record": {
            "name": "病历报告",
            "description": "电子病历文档模板",
            "sections": [
                "patient_info",
                "chief_complaint",
                "present_illness",
                "past_history",
                "physical_exam",
                "diagnosis",
                "treatment_plan"
            ]
        },
        "checkup_report": {
            "name": "体检报告",
            "description": "健康体检报告模板",
            "sections": [
                "personal_info",
                "exam_summary",
                "exam_items",
                "abnormal_findings",
                "health_suggestions"
            ]
        }
    }
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        初始化文档生成器
        
        Args:
            output_dir: 文档输出目录，默认为当前目录下的documents文件夹
        """
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            # 默认输出目录：项目根目录下的documents文件夹
            self.output_dir = Path(__file__).parent.parent.parent / "documents"
        
        # 确保输出目录存在
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 初始化模板配置
        self.templates = self.DEFAULT_TEMPLATES.copy()
        
        logger.info(f"DocumentGenerator初始化完成，输出目录: {self.output_dir}")
    
    def generate_word_document(
        self,
        title: str,
        content: Dict[str, Any],
        template: str = "health_report"
    ) -> Tuple[bytes, str]:
        """
        生成Word文档
        
        使用python-docx库生成Word文档(.docx格式)。
        
        Args:
            title: 文档标题
            content: 文档内容字典，包含各章节数据
            template: 模板名称，默认为"health_report"
            
        Returns:
            Tuple[bytes, str]: (文档字节流, 文件保存路径)
            
        Raises:
            ValueError: 当模板不存在时
            ImportError: 当python-docx未安装时
            Exception: 文档生成过程中的其他错误
            
        Example:
            >>> generator = DocumentGenerator()
            >>> content = {
            ...     "patient_name": "张三",
            ...     "age": 35,
            ...     "health_score": 85,
            ...     "recommendations": ["多运动", "少熬夜"]
            ... }
            >>> doc_bytes, file_path = generator.generate_word_document(
            ...     "健康评估报告", content, "health_report"
            ... )
        """
        logger.info(f"开始生成Word文档: title={title}, template={template}")
        
        # 检查模板是否存在
        if template not in self.templates:
            error_msg = f"模板 '{template}' 不存在，可用模板: {list(self.templates.keys())}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        try:
            # 导入python-docx
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            error_msg = "python-docx库未安装，请执行: pip install python-docx"
            logger.error(error_msg)
            raise ImportError(error_msg)
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置文档默认字体
            doc.styles['Normal'].font.name = 'Microsoft YaHei'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            # 根据模板生成文档内容
            if template == "health_report":
                self._health_report_template(doc, title, content)
            elif template == "medical_record":
                self._medical_record_template(doc, title, content)
            elif template == "checkup_report":
                self._checkup_report_template(doc, title, content)
            else:
                # 通用模板
                self._generic_template(doc, title, content)
            
            # 将文档保存到内存
            doc_bytes_io = io.BytesIO()
            doc.save(doc_bytes_io)
            doc_bytes_io.seek(0)
            doc_bytes = doc_bytes_io.getvalue()
            
            # 保存到文件
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{template}_{timestamp}.docx"
            file_path = self.save_document(doc_bytes, filename, "word")
            
            logger.info(f"Word文档生成成功: {file_path}")
            return doc_bytes, file_path
            
        except Exception as e:
            logger.error(f"生成Word文档时发生错误: {str(e)}", exc_info=True)
            raise
    
    def generate_pdf_document(
        self,
        title: str,
        content: Dict[str, Any],
        template: str = "health_report"
    ) -> Tuple[bytes, str]:
        """
        生成PDF文档
        
        使用markdown+weasyprint或docx2pdf生成PDF文档。
        优先使用weasyprint方案，如果不可用则使用Word转换方案。
        
        Args:
            title: 文档标题
            content: 文档内容字典，包含各章节数据
            template: 模板名称，默认为"health_report"
            
        Returns:
            Tuple[bytes, str]: (文档字节流, 文件保存路径)
            
        Raises:
            ValueError: 当模板不存在时
            ImportError: 当必要的PDF库未安装时
            Exception: 文档生成过程中的其他错误
            
        Example:
            >>> generator = DocumentGenerator()
            >>> content = {"patient_name": "张三", "health_score": 85}
            >>> pdf_bytes, file_path = generator.generate_pdf_document(
            ...     "健康评估报告", content, "health_report"
            ... )
        """
        logger.info(f"开始生成PDF文档: title={title}, template={template}")
        
        # 检查模板是否存在
        if template not in self.templates:
            error_msg = f"模板 '{template}' 不存在，可用模板: {list(self.templates.keys())}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        # 尝试使用weasyprint方案
        try:
            return self._generate_pdf_with_weasyprint(title, content, template)
        except ImportError:
            logger.warning("weasyprint未安装，尝试使用Word转换方案")
        except Exception as e:
            logger.warning(f"weasyprint生成失败: {str(e)}，尝试使用Word转换方案")
        
        # 回退到Word转换方案
        return self._generate_pdf_from_word(title, content, template)
    
    def _generate_pdf_with_weasyprint(
        self,
        title: str,
        content: Dict[str, Any],
        template: str
    ) -> Tuple[bytes, str]:
        """
        使用weasyprint生成PDF
        
        将内容转换为Markdown，再转为HTML，最后生成PDF。
        
        Args:
            title: 文档标题
            content: 文档内容
            template: 模板名称
            
        Returns:
            Tuple[bytes, str]: (PDF字节流, 文件路径)
        """
        try:
            import markdown
            from weasyprint import HTML, CSS
        except ImportError:
            raise ImportError("weasyprint或markdown未安装")
        
        # 生成Markdown内容
        md_content = self._content_to_markdown(title, content, template)
        
        # Markdown转HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 添加样式
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: "Microsoft YaHei", "SimHei", sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                }}
                h1 {{
                    color: #2c3e50;
                    text-align: center;
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #34495e;
                    border-left: 4px solid #3498db;
                    padding-left: 10px;
                }}
                h3 {{
                    color: #7f8c8d;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 10px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #3498db;
                    color: white;
                }}
                .highlight {{
                    background-color: #fff3cd;
                    padding: 10px;
                    border-radius: 5px;
                }}
                .footer {{
                    margin-top: 30px;
                    text-align: center;
                    font-size: 10pt;
                    color: #7f8c8d;
                }}
            </style>
        </head>
        <body>
            {html_content}
            <div class="footer">
                <p>生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
                <p>本报告由健康管理系统自动生成</p>
            </div>
        </body>
        </html>
        """
        
        # 生成PDF
        html_obj = HTML(string=styled_html)
        pdf_bytes = html_obj.write_pdf()
        
        # 保存到文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{template}_{timestamp}.pdf"
        file_path = self.save_document(pdf_bytes, filename, "pdf")
        
        logger.info(f"PDF文档生成成功(weasyprint): {file_path}")
        return pdf_bytes, file_path
    
    def _generate_pdf_from_word(
        self,
        title: str,
        content: Dict[str, Any],
        template: str
    ) -> Tuple[bytes, str]:
        """
        使用Word转换生成PDF (Windows平台)
        
        先生成Word文档，然后使用docx2pdf转换为PDF。
        
        Args:
            title: 文档标题
            content: 文档内容
            template: 模板名称
            
        Returns:
            Tuple[bytes, str]: (PDF字节流, 文件路径)
        """
        try:
            from docx2pdf import convert
        except ImportError:
            error_msg = "docx2pdf库未安装，请执行: pip install docx2pdf"
            logger.error(error_msg)
            raise ImportError(error_msg)
        
        # 先生成Word文档
        doc_bytes, doc_path = self.generate_word_document(title, content, template)
        
        # 转换PDF路径
        pdf_path = doc_path.replace(".docx", ".pdf")
        
        # 转换Word为PDF
        convert(doc_path, pdf_path)
        
        # 读取PDF字节流
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        
        logger.info(f"PDF文档生成成功(docx2pdf): {pdf_path}")
        return pdf_bytes, pdf_path
    
    def _health_report_template(
        self,
        doc: Any,
        title: str,
        content: Dict[str, Any]
    ) -> None:
        """
        健康报告模板
        
        定义健康报告的标准结构，包括基本信息、健康摘要、
        生命体征、检验结果、健康建议和随访计划等章节。
        
        Args:
            doc: python-docx Document对象
            title: 文档标题
            content: 内容字典
        """
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 文档标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行
        
        # 1. 基本信息
        doc.add_heading('一、基本信息', level=1)
        basic_info = content.get('basic_info', {})
        if basic_info:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '项目'
            hdr_cells[1].text = '内容'
            
            for key, value in basic_info.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
        else:
            # 兼容直接传入的内容
            info_items = [
                ('姓名', content.get('patient_name', 'N/A')),
                ('性别', content.get('gender', 'N/A')),
                ('年龄', str(content.get('age', 'N/A'))),
                ('身份证号', content.get('id_card', 'N/A')),
                ('联系电话', content.get('phone', 'N/A'))
            ]
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '项目'
            hdr_cells[1].text = '内容'
            
            for key, value in info_items:
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = value
        
        doc.add_paragraph()
        
        # 2. 健康摘要
        doc.add_heading('二、健康摘要', level=1)
        health_summary = content.get('health_summary', {})
        if health_summary:
            for key, value in health_summary.items():
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value))
        else:
            # 兼容直接传入的内容
            summary_text = content.get('summary', '暂无健康摘要信息')
            doc.add_paragraph(summary_text)
        
        # 健康评分
        health_score = content.get('health_score')
        if health_score is not None:
            p = doc.add_paragraph()
            p.add_run('健康评分: ').bold = True
            score_run = p.add_run(str(health_score))
            # 根据分数设置颜色
            if isinstance(health_score, (int, float)):
                if health_score >= 80:
                    score_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                elif health_score >= 60:
                    score_run.font.color.rgb = RGBColor(255, 165, 0)  # 橙色
                else:
                    score_run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        doc.add_paragraph()
        
        # 3. 生命体征
        doc.add_heading('三、生命体征', level=1)
        vital_signs = content.get('vital_signs', {})
        if vital_signs:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '指标'
            hdr_cells[1].text = '数值'
            hdr_cells[2].text = '参考范围'
            
            for key, value in vital_signs.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                if isinstance(value, dict):
                    row_cells[1].text = str(value.get('value', 'N/A'))
                    row_cells[2].text = str(value.get('reference', 'N/A'))
                else:
                    row_cells[1].text = str(value)
                    row_cells[2].text = 'N/A'
        else:
            doc.add_paragraph('暂无生命体征数据')
        
        doc.add_paragraph()
        
        # 4. 检验结果
        doc.add_heading('四、检验结果', level=1)
        lab_results = content.get('lab_results', [])
        if lab_results:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '检验项目'
            hdr_cells[1].text = '结果'
            hdr_cells[2].text = '单位'
            hdr_cells[3].text = '参考值'
            
            for item in lab_results:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('name', 'N/A'))
                result_value = str(item.get('value', 'N/A'))
                row_cells[1].text = result_value
                # 异常值标记
                if item.get('is_abnormal'):
                    row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    row_cells[1].paragraphs[0].runs[0].bold = True
                row_cells[2].text = str(item.get('unit', 'N/A'))
                row_cells[3].text = str(item.get('reference', 'N/A'))
        else:
            doc.add_paragraph('暂无检验结果数据')
        
        doc.add_paragraph()
        
        # 5. 健康建议
        doc.add_heading('五、健康建议', level=1)
        recommendations = content.get('recommendations', [])
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(rec)
        else:
            doc.add_paragraph('暂无健康建议')
        
        doc.add_paragraph()
        
        # 6. 随访计划
        doc.add_heading('六、随访计划', level=1)
        follow_up = content.get('follow_up', {})
        if follow_up:
            for key, value in follow_up.items():
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value))
        else:
            doc.add_paragraph('暂无随访计划')
        
        # 页脚
        doc.add_paragraph()
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"\n本报告由健康管理系统自动生成\n报告编号: {datetime.now().strftime('%Y%m%d%H%M%S')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _medical_record_template(
        self,
        doc: Any,
        title: str,
        content: Dict[str, Any]
    ) -> None:
        """
        病历报告模板
        
        Args:
            doc: python-docx Document对象
            title: 文档标题
            content: 内容字典
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 文档标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 各章节
        sections = [
            ('一、患者信息', 'patient_info'),
            ('二、主诉', 'chief_complaint'),
            ('三、现病史', 'present_illness'),
            ('四、既往史', 'past_history'),
            ('五、体格检查', 'physical_exam'),
            ('六、诊断', 'diagnosis'),
            ('七、治疗方案', 'treatment_plan')
        ]
        
        for section_title, key in sections:
            doc.add_heading(section_title, level=1)
            section_content = content.get(key, '暂无信息')
            
            if isinstance(section_content, dict):
                for k, v in section_content.items():
                    p = doc.add_paragraph()
                    p.add_run(f'{k}: ').bold = True
                    p.add_run(str(v))
            elif isinstance(section_content, list):
                for item in section_content:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(section_content))
            
            doc.add_paragraph()
    
    def _checkup_report_template(
        self,
        doc: Any,
        title: str,
        content: Dict[str, Any]
    ) -> None:
        """
        体检报告模板
        
        Args:
            doc: python-docx Document对象
            title: 文档标题
            content: 内容字典
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 文档标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 个人信息
        doc.add_heading('一、个人信息', level=1)
        personal_info = content.get('personal_info', {})
        if personal_info:
            for key, value in personal_info.items():
                p = doc.add_paragraph()
                p.add_run(f'{key}: ').bold = True
                p.add_run(str(value))
        
        doc.add_paragraph()
        
        # 体检摘要
        doc.add_heading('二、体检摘要', level=1)
        exam_summary = content.get('exam_summary', '暂无摘要')
        doc.add_paragraph(str(exam_summary))
        
        doc.add_paragraph()
        
        # 体检项目
        doc.add_heading('三、体检项目', level=1)
        exam_items = content.get('exam_items', [])
        if exam_items:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '检查项目'
            hdr_cells[1].text = '检查结果'
            hdr_cells[2].text = '结论'
            
            for item in exam_items:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('name', 'N/A'))
                row_cells[1].text = str(item.get('result', 'N/A'))
                row_cells[2].text = str(item.get('conclusion', 'N/A'))
        
        doc.add_paragraph()
        
        # 异常发现
        doc.add_heading('四、异常发现', level=1)
        abnormal_findings = content.get('abnormal_findings', [])
        if abnormal_findings:
            for finding in abnormal_findings:
                doc.add_paragraph(str(finding), style='List Bullet')
        else:
            doc.add_paragraph('未发现明显异常')
        
        doc.add_paragraph()
        
        # 健康建议
        doc.add_heading('五、健康建议', level=1)
        health_suggestions = content.get('health_suggestions', [])
        if health_suggestions:
            for i, suggestion in enumerate(health_suggestions, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(suggestion)
        else:
            doc.add_paragraph('请保持健康的生活方式')
    
    def _generic_template(
        self,
        doc: Any,
        title: str,
        content: Dict[str, Any]
    ) -> None:
        """
        通用模板
        
        当没有匹配到特定模板时使用。
        
        Args:
            doc: python-docx Document对象
            title: 文档标题
            content: 内容字典
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 文档标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 遍历内容字典生成文档
        for key, value in content.items():
            doc.add_heading(str(key), level=1)
            
            if isinstance(value, dict):
                for k, v in value.items():
                    p = doc.add_paragraph()
                    p.add_run(f'{k}: ').bold = True
                    p.add_run(str(v))
            elif isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(value))
            
            doc.add_paragraph()
    
    def _content_to_markdown(
        self,
        title: str,
        content: Dict[str, Any],
        template: str
    ) -> str:
        """
        将内容转换为Markdown格式
        
        Args:
            title: 文档标题
            content: 内容字典
            template: 模板名称
            
        Returns:
            str: Markdown格式的字符串
        """
        md_lines = [f"# {title}", ""]
        md_lines.append(f"**生成日期**: {datetime.now().strftime('%Y年%m月%d日')}")
        md_lines.append("")
        
        if template == "health_report":
            # 基本信息
            md_lines.append("## 一、基本信息")
            md_lines.append("")
            basic_info = content.get('basic_info', {})
            if basic_info:
                md_lines.append("| 项目 | 内容 |")
                md_lines.append("|------|------|")
                for key, value in basic_info.items():
                    md_lines.append(f"| {key} | {value} |")
            else:
                md_lines.append("| 项目 | 内容 |")
                md_lines.append("|------|------|")
                md_lines.append(f"| 姓名 | {content.get('patient_name', 'N/A')} |")
                md_lines.append(f"| 年龄 | {content.get('age', 'N/A')} |")
                md_lines.append(f"| 性别 | {content.get('gender', 'N/A')} |")
            md_lines.append("")
            
            # 健康摘要
            md_lines.append("## 二、健康摘要")
            md_lines.append("")
            health_score = content.get('health_score')
            if health_score:
                md_lines.append(f"**健康评分**: {health_score}")
            summary = content.get('summary') or content.get('health_summary', '暂无摘要')
            if isinstance(summary, dict):
                for k, v in summary.items():
                    md_lines.append(f"- **{k}**: {v}")
            else:
                md_lines.append(str(summary))
            md_lines.append("")
            
            # 生命体征
            md_lines.append("## 三、生命体征")
            md_lines.append("")
            vital_signs = content.get('vital_signs', {})
            if vital_signs:
                md_lines.append("| 指标 | 数值 | 参考范围 |")
                md_lines.append("|------|------|----------|")
                for key, value in vital_signs.items():
                    if isinstance(value, dict):
                        md_lines.append(f"| {key} | {value.get('value', 'N/A')} | {value.get('reference', 'N/A')} |")
                    else:
                        md_lines.append(f"| {key} | {value} | N/A |")
            else:
                md_lines.append("暂无数据")
            md_lines.append("")
            
            # 检验结果
            md_lines.append("## 四、检验结果")
            md_lines.append("")
            lab_results = content.get('lab_results', [])
            if lab_results:
                md_lines.append("| 检验项目 | 结果 | 单位 | 参考值 |")
                md_lines.append("|----------|------|------|--------|")
                for item in lab_results:
                    result = str(item.get('value', 'N/A'))
                    if item.get('is_abnormal'):
                        result += " ⚠️"
                    md_lines.append(
                        f"| {item.get('name', 'N/A')} | {result} | "
                        f"{item.get('unit', 'N/A')} | {item.get('reference', 'N/A')} |"
                    )
            else:
                md_lines.append("暂无数据")
            md_lines.append("")
            
            # 健康建议
            md_lines.append("## 五、健康建议")
            md_lines.append("")
            recommendations = content.get('recommendations', [])
            if recommendations:
                for i, rec in enumerate(recommendations, 1):
                    md_lines.append(f"{i}. {rec}")
            else:
                md_lines.append("暂无建议")
            md_lines.append("")
            
            # 随访计划
            md_lines.append("## 六、随访计划")
            md_lines.append("")
            follow_up = content.get('follow_up', {})
            if follow_up:
                for k, v in follow_up.items():
                    md_lines.append(f"- **{k}**: {v}")
            else:
                md_lines.append("暂无计划")
        else:
            # 通用Markdown转换
            for key, value in content.items():
                md_lines.append(f"## {key}")
                md_lines.append("")
                if isinstance(value, dict):
                    for k, v in value.items():
                        md_lines.append(f"- **{k}**: {v}")
                elif isinstance(value, list):
                    for item in value:
                        md_lines.append(f"- {item}")
                else:
                    md_lines.append(str(value))
                md_lines.append("")
        
        return "\n".join(md_lines)
    
    def save_document(
        self,
        doc_bytes: bytes,
        filename: str,
        doc_type: str
    ) -> str:
        """
        保存文档到本地
        
        Args:
            doc_bytes: 文档字节流
            filename: 文件名
            doc_type: 文档类型 (word/pdf)
            
        Returns:
            str: 保存的文件完整路径
            
        Raises:
            ValueError: 当文档类型不支持时
            IOError: 当文件保存失败时
        """
        logger.info(f"保存文档: filename={filename}, type={doc_type}")
        
        # 验证文档类型
        if doc_type.lower() not in self.SUPPORTED_DOC_TYPES:
            error_msg = f"不支持的文档类型: {doc_type}，支持类型: {self.SUPPORTED_DOC_TYPES}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        try:
            # 根据类型选择子目录
            type_dir = self.output_dir / doc_type.lower()
            type_dir.mkdir(parents=True, exist_ok=True)
            
            # 构建完整文件路径
            file_path = type_dir / filename
            
            # 写入文件
            with open(file_path, 'wb') as f:
                f.write(doc_bytes)
            
            logger.info(f"文档保存成功: {file_path}")
            return str(file_path)
            
        except IOError as e:
            error_msg = f"保存文档失败: {str(e)}"
            logger.error(error_msg)
            raise IOError(error_msg)
    
    def get_available_templates(self) -> Dict[str, Dict[str, str]]:
        """
        获取可用模板列表
        
        Returns:
            Dict: 模板名称到模板信息的映射
        """
        return {
            name: {
                "name": info["name"],
                "description": info["description"]
            }
            for name, info in self.templates.items()
        }
    
    def add_custom_template(
        self,
        template_name: str,
        template_config: Dict[str, Any]
    ) -> None:
        """
        添加自定义模板
        
        Args:
            template_name: 模板名称
            template_config: 模板配置字典
        """
        self.templates[template_name] = template_config
        logger.info(f"添加自定义模板: {template_name}")


# 便捷函数，用于快速生成文档
def generate_health_report(
    patient_name: str,
    health_data: Dict[str, Any],
    output_format: str = "word",
    output_dir: Optional[str] = None
) -> Tuple[bytes, str]:
    """
    快速生成健康报告
    
    便捷函数，用于快速生成健康报告文档。
    
    Args:
        patient_name: 患者姓名
        health_data: 健康数据字典
        output_format: 输出格式 (word/pdf)
        output_dir: 输出目录
        
    Returns:
        Tuple[bytes, str]: (文档字节流, 文件路径)
    """
    generator = DocumentGenerator(output_dir)
    
    content = {
        "patient_name": patient_name,
        **health_data
    }
    
    title = f"{patient_name}的健康评估报告"
    
    if output_format.lower() == "pdf":
        return generator.generate_pdf_document(title, content, "health_report")
    else:
        return generator.generate_word_document(title, content, "health_report")


# 模块测试代码
if __name__ == "__main__":
    # 配置日志
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # 测试数据
    test_content = {
        "basic_info": {
            "姓名": "张三",
            "性别": "男",
            "年龄": "35岁",
            "身份证号": "110101199001011234",
            "联系电话": "13800138000"
        },
        "health_summary": {
            "总体评价": "健康状况良好",
            "主要风险": "轻度脂肪肝",
            "建议": "控制饮食，增加运动"
        },
        "health_score": 85,
        "vital_signs": {
            "血压": {"value": "120/80 mmHg", "reference": "90-140/60-90 mmHg"},
            "心率": {"value": "72 次/分", "reference": "60-100 次/分"},
            "体温": {"value": "36.5°C", "reference": "36.0-37.2°C"},
            "体重": {"value": "70 kg", "reference": "N/A"},
            "身高": {"value": "175 cm", "reference": "N/A"},
            "BMI": {"value": "22.9", "reference": "18.5-23.9"}
        },
        "lab_results": [
            {"name": "血红蛋白", "value": "145", "unit": "g/L", "reference": "130-175", "is_abnormal": False},
            {"name": "白细胞", "value": "6.5", "unit": "10^9/L", "reference": "4.0-10.0", "is_abnormal": False},
            {"name": "血糖", "value": "5.8", "unit": "mmol/L", "reference": "3.9-6.1", "is_abnormal": False},
            {"name": "总胆固醇", "value": "5.8", "unit": "mmol/L", "reference": "<5.2", "is_abnormal": True},
            {"name": "甘油三酯", "value": "2.1", "unit": "mmol/L", "reference": "<1.7", "is_abnormal": True}
        ],
        "recommendations": [
            "减少高脂肪食物摄入，多吃蔬菜水果",
            "每周进行至少150分钟中等强度有氧运动",
            "保持规律作息，避免熬夜",
            "定期复查血脂指标，建议3个月后复查",
            "控制体重，BMI保持在正常范围"
        ],
        "follow_up": {
            "下次随访时间": "2026年6月25日",
            "随访项目": "血脂四项、肝功能",
            "注意事项": "空腹检查"
        }
    }
    
    # 创建文档生成器
    generator = DocumentGenerator()
    
    # 显示可用模板
    print("可用模板:")
    for name, info in generator.get_available_templates().items():
        print(f"  - {name}: {info['description']}")
    
    # 生成Word文档
    print("\n正在生成Word文档...")
    try:
        doc_bytes, doc_path = generator.generate_word_document(
            "健康评估报告",
            test_content,
            "health_report"
        )
        print(f"Word文档生成成功!")
        print(f"文件路径: {doc_path}")
        print(f"文件大小: {len(doc_bytes)} bytes")
    except Exception as e:
        print(f"Word文档生成失败: {str(e)}")
    
    # 生成PDF文档
    print("\n正在生成PDF文档...")
    try:
        pdf_bytes, pdf_path = generator.generate_pdf_document(
            "健康评估报告",
            test_content,
            "health_report"
        )
        print(f"PDF文档生成成功!")
        print(f"文件路径: {pdf_path}")
        print(f"文件大小: {len(pdf_bytes)} bytes")
    except Exception as e:
        print(f"PDF文档生成失败: {str(e)}")
        print("提示: 请安装必要的依赖库:")
        print("  - weasyprint: pip install weasyprint")
        print("  - docx2pdf: pip install docx2pdf (Windows)")
